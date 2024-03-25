import os

from docx import Document
from docx.shared import Inches

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def read_figures_from_directory(directory):
    # Get all files in the directory
    files = os.listdir(directory)

    # Filter the files to include only image files
    image_files = [file for file in files if file.endswith(('.png', '.jpg', '.jpeg'))]
    
    # Sort the image files based on their filenames
    image_files = sorted(image_files)


    # Create a list to store the full paths of the image files
    image_paths = []

    # Iterate over the image files and create the full paths
    for file in image_files:
        image_path = os.path.join(directory, file)
        image_paths.append(image_path)

    return image_paths


def create_word_file(figures, word_file):
    # Create a new Word document
    doc = Document()

    # Add a two-column table to the document
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Iterate over the figures and insert them into the table
    for i, figure in enumerate(figures):
        # Add a new row to the table for each figure
        if i % 2 == 0:
            row_cells = table.add_row().cells

        # Calculate the cell index based on the figure index
        cell_index = i % 2

        # Add the figure to the cell
        cell = row_cells[cell_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(figure, width=Inches(3), height=Inches(3))

    # Save the Word document
    doc.save(word_file)

